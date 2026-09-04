"""Unit tests for multi-format DocumentParser service (Phase 4.2.a)."""

import io

import docx
import pypdf
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject
import pytest

from plutolab_api.services.doc_parser import (
    DocParseError,
    DocumentParser,
    ParsedDocument,
    ParsedPage,
)


def _create_sample_pdf(pages_text: list[str]) -> bytes:
    """Helper to create an in-memory PDF with extractable text streams."""
    writer = pypdf.PdfWriter()

    fonts = DictionaryObject()
    font = DictionaryObject()
    font[NameObject("/Type")] = NameObject("/Font")
    font[NameObject("/Subtype")] = NameObject("/Type1")
    font[NameObject("/BaseFont")] = NameObject("/Helvetica")
    fonts[NameObject("/F1")] = font

    resources = DictionaryObject()
    resources[NameObject("/Font")] = fonts

    for text in pages_text:
        page = writer.add_blank_page(width=300, height=300)
        page[NameObject("/Resources")] = resources
        stream = DecodedStreamObject()
        # PDF text operator: BT /F1 12 Tf (text) Tj ET
        escaped = text.replace("(", r"\(").replace(")", r"\)")
        stream.set_data(f"BT /F1 12 Tf 50 200 Td ({escaped}) Tj ET".encode("latin-1"))
        page[NameObject("/Contents")] = stream

    bio = io.BytesIO()
    writer.write(bio)
    return bio.getvalue()


def _create_sample_docx(
    paragraphs: list[str], table_data: list[list[str]] | None = None
) -> bytes:
    """Helper to create an in-memory Word document."""
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)

    if table_data:
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
        table = doc.add_table(rows=rows, cols=cols)
        for r_idx, row in enumerate(table_data):
            for c_idx, val in enumerate(row):
                table.cell(r_idx, c_idx).text = val

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def test_parse_plain_text_utf8() -> None:
    content = "Hello, PlutoLab!\n这是中文测试内容。\n".encode("utf-8")
    parsed = DocumentParser.parse(content, "test.txt")

    assert isinstance(parsed, ParsedDocument)
    assert parsed.file_type == "txt"
    assert "Hello, PlutoLab!" in parsed.text
    assert "这是中文测试内容。" in parsed.text
    assert parsed.page_count == 1
    assert len(parsed.pages) == 1
    assert parsed.pages[0].page_number == 1
    assert parsed.char_count == len(parsed.text)


def test_parse_plain_text_gbk_fallback() -> None:
    # Text encoded with Chinese GBK encoding
    chinese_text = "这是使用 GBK 编码保存的简体中文文档。"
    content = chinese_text.encode("gbk")

    parsed = DocumentParser.parse(content, "notes_gbk.txt")
    assert parsed.file_type == "txt"
    assert parsed.text == chinese_text
    assert parsed.char_count == len(chinese_text)


def test_parse_markdown_basic_and_frontmatter() -> None:
    md_content = (
        "---\n"
        "title: RAG 架构设计\n"
        "author: Pluto\n"
        "---\n\n"
        "# 核心架构\n\n"
        "本方案采用 pgvector 进行向量存储，结合全文检索实现混合召回。"
    ).encode("utf-8")

    parsed = DocumentParser.parse(md_content, "architecture.md")

    assert parsed.file_type == "md"
    assert "# 核心架构" in parsed.text
    assert "pgvector" in parsed.text
    assert "frontmatter_raw" in parsed.metadata
    assert "title: RAG 架构设计" in parsed.metadata["frontmatter_raw"]


def test_parse_pdf_single_and_multi_page() -> None:
    pdf_bytes = _create_sample_pdf(
        ["Page 1: Introduction to AI Agents", "Page 2: Retrieval Augmented Generation"]
    )
    parsed = DocumentParser.parse(pdf_bytes, "agents.pdf")

    assert parsed.file_type == "pdf"
    assert parsed.page_count == 2
    assert len(parsed.pages) == 2
    assert parsed.pages[0].page_number == 1
    assert "Page 1: Introduction to AI Agents" in parsed.pages[0].text
    assert parsed.pages[1].page_number == 2
    assert "Page 2: Retrieval Augmented Generation" in parsed.pages[1].text
    assert parsed.metadata["total_physical_pages"] == 2


def test_parse_docx_paragraphs_and_tables() -> None:
    docx_bytes = _create_sample_docx(
        paragraphs=["Project PlutoLab Overview", "Key Milestones achieved in 2026."],
        table_data=[["Phase", "Status"], ["Phase 4.1", "Completed"], ["Phase 4.2", "In Progress"]],
    )
    parsed = DocumentParser.parse(docx_bytes, "milestones.docx")

    assert parsed.file_type == "docx"
    assert "Project PlutoLab Overview" in parsed.text
    assert "Phase | Status" in parsed.text
    assert "Phase 4.1 | Completed" in parsed.text
    assert parsed.page_count == 1


def test_parse_empty_content_raises_error() -> None:
    with pytest.raises(DocParseError, match="Document is empty"):
        DocumentParser.parse(b"", "empty.txt")

    with pytest.raises(DocParseError, match="Document is empty"):
        DocumentParser.parse(b"   \n\t  ", "blank.txt")


def test_parse_unsupported_extension_raises_error() -> None:
    with pytest.raises(DocParseError, match="Unsupported file format '.exe'"):
        DocumentParser.parse(b"dummy binary", "malicious.exe")


def test_parse_corrupted_pdf_raises_error() -> None:
    corrupted = b"%PDF-1.4\nCorrupted content without valid xref or trailer"
    with pytest.raises(DocParseError, match="Failed to open PDF document"):
        DocumentParser.parse(corrupted, "corrupted.pdf")


def test_parse_corrupted_docx_raises_error() -> None:
    corrupted = b"PK\x03\x04not a valid zip package for docx"
    with pytest.raises(DocParseError, match="Failed to open Word"):
        DocumentParser.parse(corrupted, "corrupted.docx")


def test_parse_file_size_exceeded_raises_error() -> None:
    content = b"x" * 1024
    with pytest.raises(DocParseError, match="File size exceeds maximum allowed limit"):
        DocumentParser.parse(content, "large.txt", max_size=512)


def test_parse_scanned_pdf_without_text_raises_error() -> None:
    # Blank PDF with no text stream
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=100, height=100)
    bio = io.BytesIO()
    writer.write(bio)

    with pytest.raises(DocParseError, match="PDF contains no extractable text"):
        DocumentParser.parse(bio.getvalue(), "scanned_image.pdf")
