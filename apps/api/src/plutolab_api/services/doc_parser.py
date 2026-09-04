"""Multi-format document parser service for RAG ingestion (Phase 4.2.a).

Supports:
- Markdown (.md, .markdown)
- Plain text (.txt)
- PDF (.pdf) via pypdf
- Word (.docx) via python-docx
"""

from dataclasses import dataclass, field
import io
from pathlib import Path
from typing import Any

import docx
from docx.opc.exceptions import PackageNotFoundError
import pypdf
from pypdf.errors import PdfReadError
import structlog

logger = structlog.get_logger(__name__)

# Max file size limit: 50MB
MAX_FILE_SIZE = 50 * 1024 * 1024

# Encodings to try sequentially for text/markdown
TEXT_ENCODINGS = ["utf-8", "utf-8-sig", "gb18030", "gbk", "latin-1"]


class DocParseError(Exception):
    """Domain exception raised when document parsing fails."""

    def __init__(self, message: str, original_error: Exception | None = None) -> None:
        super().__init__(message)
        self.original_error = original_error


@dataclass
class ParsedPage:
    """Individual parsed page / section with physical or logical sequence index."""

    page_number: int  # 1-indexed
    text: str
    char_count: int


@dataclass
class ParsedDocument:
    """Standardized output of the document parser."""

    text: str
    char_count: int
    page_count: int
    pages: list[ParsedPage]
    metadata: dict[str, Any] = field(default_factory=dict)
    file_type: str = "txt"


class DocumentParser:
    """High-throughput in-memory multi-format document parser."""

    @classmethod
    def parse(
        cls,
        content: bytes,
        filename: str,
        file_type: str | None = None,
        max_size: int = MAX_FILE_SIZE,
    ) -> ParsedDocument:
        """Parse raw document bytes into a structured `ParsedDocument`.

        Args:
            content: Raw document binary data
            filename: Original file name (used for extension detection)
            file_type: Explicit file extension override (e.g. "pdf", "md")
            max_size: Maximum allowed byte size (default 50MB)

        Returns:
            ParsedDocument with extracted pages, text, and metadata.

        Raises:
            DocParseError: If file is too large, empty, corrupted, or unsupported.
        """
        if not content or len(content.strip()) == 0:
            raise DocParseError("Document is empty or contains only whitespace bytes.")

        if len(content) > max_size:
            mb = max_size // (1024 * 1024)
            raise DocParseError(f"File size exceeds maximum allowed limit of {mb}MB.")

        # Determine normalized extension / file type
        ext = file_type or Path(filename).suffix.lstrip(".").lower()
        if not ext:
            ext = "txt"

        if ext in ("md", "markdown"):
            return cls._parse_markdown(content, filename)
        elif ext == "txt":
            return cls._parse_text(content, filename)
        elif ext == "pdf":
            return cls._parse_pdf(content, filename)
        elif ext == "docx":
            return cls._parse_docx(content, filename)
        else:
            raise DocParseError(
                f"Unsupported file format '.{ext}'. Supported formats: .md, .txt, .pdf, .docx"
            )

    @classmethod
    def _decode_text_bytes(cls, content: bytes) -> str:
        """Decode text bytes trying multiple character encodings."""
        for enc in TEXT_ENCODINGS:
            try:
                return content.decode(enc)
            except (UnicodeDecodeError, LookupError):
                continue
        raise DocParseError(
            "Unable to decode text document. Unsupported text character encoding."
        )

    @classmethod
    def _parse_text(cls, content: bytes, filename: str) -> ParsedDocument:
        """Parse plain text files (.txt)."""
        raw_text = cls._decode_text_bytes(content).strip()
        if not raw_text:
            raise DocParseError("Text document contains no readable characters.")

        char_count = len(raw_text)
        page = ParsedPage(page_number=1, text=raw_text, char_count=char_count)

        return ParsedDocument(
            text=raw_text,
            char_count=char_count,
            page_count=1,
            pages=[page],
            metadata={"filename": filename, "format": "txt"},
            file_type="txt",
        )

    @classmethod
    def _parse_markdown(cls, content: bytes, filename: str) -> ParsedDocument:
        """Parse Markdown files (.md), preserving markdown structure."""
        raw_text = cls._decode_text_bytes(content).strip()
        if not raw_text:
            raise DocParseError("Markdown document contains no readable characters.")

        metadata: dict[str, Any] = {"filename": filename, "format": "md"}

        # Basic frontmatter parsing if present (--- ... ---)
        body_text = raw_text
        if raw_text.startswith("---"):
            parts = raw_text.split("---", 2)
            if len(parts) >= 3:
                frontmatter_raw = parts[1].strip()
                body_text = parts[2].strip()
                metadata["frontmatter_raw"] = frontmatter_raw

        char_count = len(body_text) if body_text else len(raw_text)
        actual_text = body_text if body_text else raw_text

        page = ParsedPage(page_number=1, text=actual_text, char_count=char_count)

        return ParsedDocument(
            text=actual_text,
            char_count=char_count,
            page_count=1,
            pages=[page],
            metadata=metadata,
            file_type="md",
        )

    @classmethod
    def _parse_pdf(cls, content: bytes, filename: str) -> ParsedDocument:
        """Parse PDF files (.pdf) using pypdf."""
        try:
            stream = io.BytesIO(content)
            reader = pypdf.PdfReader(stream)
        except Exception as exc:
            raise DocParseError("Failed to open PDF document. File may be corrupted.", exc) from exc

        if reader.is_encrypted:
            try:
                # Attempt decrypt with empty password for standard protected PDFs
                decrypted = reader.decrypt("")
                if decrypted == 0:
                    raise DocParseError("PDF document is encrypted and requires a password.")
            except Exception as exc:
                raise DocParseError(
                    "PDF document is password-protected and cannot be decrypted.", exc
                ) from exc

        if not reader.pages or len(reader.pages) == 0:
            raise DocParseError("PDF document contains no pages.")

        pages: list[ParsedPage] = []
        page_texts: list[str] = []

        for idx, page in enumerate(reader.pages):
            try:
                extracted = page.extract_text() or ""
            except Exception as exc:
                logger.warning("pdf_page_extraction_warning", page=idx + 1, error=str(exc))
                extracted = ""

            cleaned = extracted.strip()
            if cleaned:
                pages.append(
                    ParsedPage(page_number=idx + 1, text=cleaned, char_count=len(cleaned))
                )
                page_texts.append(cleaned)

        if not pages:
            raise DocParseError(
                "PDF contains no extractable text (document may be a scanned image or empty)."
            )

        full_text = "\n\n".join(page_texts)
        total_chars = sum(p.char_count for p in pages)

        meta: dict[str, Any] = {
            "filename": filename,
            "format": "pdf",
            "total_physical_pages": len(reader.pages),
        }
        if reader.metadata:
            for k in ("/Title", "/Author", "/Subject", "/Creator"):
                if k in reader.metadata and reader.metadata[k]:
                    clean_key = k.lstrip("/").lower()
                    meta[clean_key] = str(reader.metadata[k])

        return ParsedDocument(
            text=full_text,
            char_count=total_chars,
            page_count=len(pages),
            pages=pages,
            metadata=meta,
            file_type="pdf",
        )

    @classmethod
    def _parse_docx(cls, content: bytes, filename: str) -> ParsedDocument:
        """Parse Word documents (.docx) using python-docx."""
        try:
            stream = io.BytesIO(content)
            doc = docx.Document(stream)
        except (PackageNotFoundError, Exception) as exc:
            raise DocParseError(
                "Failed to open Word (.docx) document. File may be corrupted or not in OOXML format.",
                exc,
            ) from exc

        sections_text: list[str] = []

        # 1. Paragraphs
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                sections_text.append(txt)

        # 2. Tables
        for table in doc.tables:
            table_lines: list[str] = []
            for row in table.rows:
                row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_cells:
                    table_lines.append(" | ".join(row_cells))
            if table_lines:
                sections_text.append("\n".join(table_lines))

        if not sections_text:
            raise DocParseError("Word document contains no readable text or content.")

        full_text = "\n\n".join(sections_text)
        total_chars = len(full_text)

        meta: dict[str, Any] = {"filename": filename, "format": "docx"}
        try:
            props = doc.core_properties
            if props.title:
                meta["title"] = props.title
            if props.author:
                meta["author"] = props.author
        except Exception:
            pass

        page = ParsedPage(page_number=1, text=full_text, char_count=total_chars)

        return ParsedDocument(
            text=full_text,
            char_count=total_chars,
            page_count=1,
            pages=[page],
            metadata=meta,
            file_type="docx",
        )
